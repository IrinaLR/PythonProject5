import sqlalchemy
from flask import Flask, render_template, request, redirect, url_for, flash, Response, make_response, send_file
from sqlalchemy.engine import URL
from sqlalchemy.orm import Session, sessionmaker
from datetime import datetime
import pandas
import io
from io import BytesIO
import docx
import time
import os
from flask_sqlalchemy import SQLAlchemy
from sqlalchemy import create_engine
from sqlalchemy import MetaData, Table, String, Integer, Column, Text, DateTime, Boolean
import pyodbc


app = Flask(__name__)

# db connect
driver = '{SQL Server}'
server = '09-36022\SQLEXPRESS'
database = 'TestDB'
trusted_connection = 'yes'
connection_string = f'DRIVER={driver};SERVER={server};'
connection_string += f'DATABASE={database};'
connection_string += f'TRUSTED_CONNECTION={trusted_connection}'
connection_url = URL.create("mssql+pyodbc", query={"odbc_connect": connection_string})
engine = sqlalchemy.create_engine(connection_url)
metadata = sqlalchemy.MetaData(bind=engine)
session = Session(bind=engine)

# link for tables
products = sqlalchemy.Table('Products', metadata, autoload=True)

# global var
select_product = None
filter_tmn = None


@app.route('/')
@app.route('/home')  # main page
def home():
    return render_template(
        'index.html',
        title='Home',
        year=datetime.now().year,
    )


@app.route('/about')  # about vertex
def about():
    return render_template(
        'about.html',
        title='About',
        year=datetime.now().year,
        message='Hello !'
    )


@app.route('/db')  # table Products
def db():
    result_all = [list(i) for i in session.query(products).all()]
    return render_template(
        'db.html',
        title='About',
        year=datetime.now().year,
        message='Hello !', result_all=result_all
    )


@app.route('/add_data', methods=['POST'])  # insert and redirect after changes
def add_data():
    product = request.form['product']
    price = request.form['price']
    quantity = request.form['quantity']  # form for new column
    query = products.insert({'ProductName': product, 'Price': price, 'Quantity': quantity})  # + name for new column
    session.execute(query)
    session.commit()
    return redirect(url_for('db'))


@app.route('/sdb', methods=['GET', 'POST'])  # show all and select 1 position from Products
def sdb():
    results = [list(i) for i in session.query(products).all()]
    global select_product
    if request.method == 'POST':
        select_product = str(request.form['select_product'])
    price_if = None
    product_if = None
    if select_product is not None:
        results_if = [list(i) for i in session.query(products).filter_by(ProductName=select_product)]
        price_if = int(results_if[0][2])
        product_if = results_if[0][1]
    else:
        price_if = ''
        product_if = ''
    return render_template('sdb.html', results=results, price_if=price_if, year=datetime.now().year, product_if=product_if)


@app.route('/upd', methods=['GET', 'POST'])  # price change (update)
def upd():
    if request.method == 'POST':
        price2 = int(request.form['price2'])
        session.query(products).filter_by(ProductName=select_product).update({"Price": price2}, synchronize_session='false')
        session.commit()
    return render_template('upd.html', year=datetime.now().year)


@app.route('/xls')  # export to xlsx
def xls():
    out = io.BytesIO()
    df = pandas.read_sql_table('Products', con=engine)
    df.to_excel(out)
    return Response(out.getvalue(), mimetype="application/ms-excel", headers={"Content-Disposition": "attachment;filename=report.xlsx"})


@app.route('/doc')  # export to docx
def doc():
    docs = docx.Document()
    text = 'Данные из таблицы Products'
    par = docs.add_paragraph(text)
    df = pandas.read_sql_table('Products', con=engine)
    t = docs.add_table(df.shape[0] + 1, df.shape[1])
    for j in range(df.shape[-1]):
        t.cell(0, j).text = df.columns[j]
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            t.cell(i + 1, j).text = str(df.values[i, j])
    f = BytesIO()
    docs.save(f)
    f.seek(0)
    return send_file(f, as_attachment=True, download_name='report.docx')


@app.route('/tmn')  # test table 26m lines
def tmn():
    tm = sqlalchemy.Table('Datatmn', metadata, autoload=True)
    result_tmn = []
    ts = time.time()
    if filter_tmn is not None:
        result_tmn = [list(i) for i in session.query(tm).filter_by(value=filter_tmn)]
    tf = time.time() - ts
    counter = 0
    for i in result_tmn:
        i = float(i[5].replace(',', '.'))
        # i = int(i[5])
        counter += i
    return render_template(
        'tmn.html',
        title='About',
        year=datetime.now().year,
        message='Hello !', result_tmn=result_tmn, counter=counter, tf=tf
    )


@app.route('/fl', methods=['POST'])  # filter
def fl():
    global filter_tmn
    if request.method == 'POST':
        filter_tmn = str(request.form['filter_tmn'])
    return redirect(url_for('tmn'))


@app.route('/form')
def form():
    return render_template(
        'form.html',
        title='Form',
        year=datetime.now().year,
        message='Welcome !'
    )


@app.route('/sending', methods=['GET', 'POST'])
def sending():
    inp_name = request.form.get("username")
    inp_pass = request.form.get("password")
    if request.method == 'POST' and inp_name and inp_pass:
        # write to file
        with open('test.txt', 'a') as f:
            f.write(str('Username   ' + inp_name + '   Password   ' + inp_pass + '\n'))
        return render_template('sending.html', year=datetime.now().year, username=inp_name, password=inp_pass)
    return render_template(
        'sending.html',
        title='Form',
        year=datetime.now().year
    )


@app.route('/output')
def output():
    names = 'This is my variable on Python !'
    return render_template(
        'output.html',
        title='Form',
        year=datetime.now().year,
        names=names
    )


@app.route('/cat')
def cat():
    return render_template(
        'cat.html',
        title='About',
        year=datetime.now().year,
        message='Hello !'
    )


if __name__ == "__main__":
    app.run(debug=True)  # (debug=False) for user


